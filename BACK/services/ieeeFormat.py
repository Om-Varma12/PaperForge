from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --------------------------------------------------
# Roman numeral helper
# --------------------------------------------------
def to_roman(num):
    romans = {
        1: "I", 2: "II", 3: "III", 4: "IV", 5: "V",
        6: "VI", 7: "VII", 8: "VIII", 9: "IX", 10: "X"
    }
    return romans.get(num, str(num))


# --------------------------------------------------
# Continuous section (NO page break)
# --------------------------------------------------
def add_continuous_section(doc):
    section = doc.add_section()
    sectPr = section._sectPr

    type_el = sectPr.xpath("./w:type")
    if not type_el:
        type_el = OxmlElement("w:type")
        sectPr.insert(0, type_el)
    else:
        type_el = type_el[0]

    type_el.set(qn("w:val"), "continuous")
    return section


# --------------------------------------------------
# Main IEEE Paper Generator
# --------------------------------------------------
def generate_ieee_paper(paper: dict):

    def set_two_columns(section):
        sectPr = section._sectPr
        cols = sectPr.xpath("./w:cols")
        if cols:
            cols = cols[0]
        else:
            cols = OxmlElement("w:cols")
            sectPr.append(cols)

        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "720")

    # -------- Section Heading (10pt, Bold, Centered) --------
    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True

        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2.6)

    # -------- Abstract & Keywords Body (9pt, Bold) --------
    def add_bold_body_paragraph(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.bold = True

        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2.6)

    # -------- Normal Body Paragraph (9pt, Normal) --------
    def add_normal_body_paragraph(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.bold = False

        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2.6)

    # ---------------- DOCUMENT SETUP ----------------

    doc = Document()

    # 🔥 NARROW MARGINS
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ---------------- TITLE ----------------

    title = doc.add_paragraph()
    run = title.add_run(paper["title"])
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------- AUTHORS ----------------

    authors = doc.add_paragraph()
    run = authors.add_run(paper["authors"])
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.bold = False
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------- TWO COLUMN CONTINUOUS ----------------

    section = add_continuous_section(doc)
    set_two_columns(section)

    # ---------------- ABSTRACT ----------------

    add_section_heading("Abstract")
    add_bold_body_paragraph(paper["abstract"])

    # ---------------- KEYWORDS ----------------

    add_section_heading("Keywords")
    add_bold_body_paragraph(", ".join(paper["keywords"]))
    doc.add_paragraph("")

    # ---------------- MAIN SECTIONS ----------------

    for idx, sec in enumerate(paper["sections"], start=1):
        roman = to_roman(idx)
        add_section_heading(f"{roman}. {sec['title']}")
        add_normal_body_paragraph(sec["content"])

        # One empty line between sections
        doc.add_paragraph("")

    # ---------------- REFERENCES ----------------

    add_section_heading("References")
    for ref in paper["references"]:
        add_normal_body_paragraph(ref)

    # ---------------- SAVE ----------------

    doc.save(paper["output_file"])
    print(f"✅ IEEE paper generated with FINAL formatting lock 🔒")

# ======================================================================
# PAPER DATA
# ======================================================================

paper_data = {
    "title": "AI-Based Phishing Website Detection Using Machine Learning Techniques",
    "authors": (
        "Om Varma\n"
        "Department of Computer Science\n"
        "India\n"
        "omvarma@email.com"
    ),
    "abstract": (
        "Phishing attacks remain one of the most critical cybersecurity threats, "
        "exploiting user trust to steal sensitive information such as login credentials, "
        "financial data, and personal identity details. Traditional blacklist-based and "
        "rule-based detection mechanisms fail to detect newly emerging phishing websites, "
        "making them ineffective against zero-day attacks. This paper proposes an "
        "artificial intelligence driven phishing website detection system that leverages "
        "machine learning techniques to classify websites as benign or malicious. The "
        "proposed system extracts URL-based, domain-based, and content-based features, "
        "including lexical patterns, domain age, SSL certificate information, and HTML "
        "structure. An XGBoost classifier is trained on a large-scale dataset containing "
        "over six hundred thousand URLs. Experimental results demonstrate that the "
        "proposed approach achieves superior accuracy, precision, recall, and F1-score "
        "compared to traditional machine learning models, highlighting its effectiveness "
        "in real-world phishing detection scenarios."
    ),
    "keywords": [
        "Phishing Detection",
        "Machine Learning",
        "Cybersecurity",
        "XGBoost",
        "Website Security"
    ],
    "sections": [
        {
            "title": "1. Introduction",
            "content": (
                "Phishing websites are fraudulent online platforms designed to impersonate "
                "legitimate websites with the goal of deceiving users into revealing "
                "confidential information. With the rapid growth of internet usage and "
                "online services, phishing attacks have become increasingly sophisticated "
                "and difficult to detect. Attackers frequently modify URLs, clone website "
                "interfaces, and deploy short-lived domains to bypass traditional detection "
                "systems. As a result, phishing continues to pose a severe threat to both "
                "individual users and organizations worldwide. This section discusses the "
                "motivation behind phishing detection, the challenges faced by existing "
                "approaches, and the necessity for intelligent, data-driven solutions."
            )
        },
        {
            "title": "2. Related Work",
            "content": (
                "Several studies have explored phishing website detection using blacklist-based, "
                "heuristic-based, and machine learning-based approaches. Blacklist-based methods "
                "rely on previously reported phishing URLs but fail to identify newly created "
                "malicious websites. Heuristic-based approaches utilize manually crafted rules, "
                "which are often brittle and require frequent updates. Recent research has "
                "demonstrated that machine learning models such as Support Vector Machines, "
                "Random Forests, and Neural Networks can effectively classify phishing websites "
                "using extracted features. However, these models often suffer from high false "
                "positive rates or lack interpretability, limiting their practical adoption."
            )
        },
        {
            "title": "3. Methodology",
            "content": (
                "The proposed phishing detection system consists of four major components: data "
                "collection, feature extraction, model training, and prediction. A large-scale "
                "dataset containing benign, phishing, malware, and defacement URLs is collected "
                "from publicly available sources. Feature extraction involves analyzing URL "
                "lexical properties, domain registration details, SSL certificate attributes, "
                "and webpage content. An XGBoost classifier is trained using these features due "
                "to its ability to handle high-dimensional data and capture complex feature "
                "interactions. The trained model is then integrated into a real-time detection "
                "pipeline capable of analyzing URLs during browsing sessions."
            )
        },
        {
            "title": "4. Experimental Results",
            "content": (
                "Extensive experiments are conducted to evaluate the performance of the proposed "
                "model. The dataset is divided into training and testing subsets, and multiple "
                "evaluation metrics including accuracy, precision, recall, F1-score, and "
                "confusion matrix are computed. The XGBoost model achieves an accuracy exceeding "
                "99 percent, with significantly lower false positive rates compared to baseline "
                "models. These results demonstrate the robustness and generalization capability "
                "of the proposed system across diverse phishing scenarios."
            )
        },
        {
            "title": "5. Conclusion",
            "content": (
                "This paper presents an effective AI-based phishing website detection system "
                "capable of identifying both known and zero-day phishing attacks. By leveraging "
                "machine learning techniques and comprehensive feature extraction, the proposed "
                "approach significantly outperforms traditional detection mechanisms. Future "
                "work will focus on improving model explainability, reducing computational "
                "overhead, and deploying the system as a browser extension for real-time user "
                "protection."
            )
        }
    ],
    "references": [
        "[1] R. Verma and K. Dyer, \"On the Character of Phishing URLs,\" IEEE, 2015.",
        "[2] M. Aburrous et al., \"Intelligent phishing detection system,\" Expert Systems, 2010.",
        "[3] T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" KDD, 2016."
    ],
    "output_file": "IEEE_Paper_PaperForge.docx"
}
# ---------------- RUN ----------------
generate_ieee_paper(paper_data)
